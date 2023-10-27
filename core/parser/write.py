from docx.shared import RGBColor
from docx import Document


async def create_document(_token: str, _chat_id: str):
    _document = Document()
    _document.add_heading("Отчет по тестированию бота: " + _token)
    _document.add_heading("Id чата: " + _chat_id)
    _document.add_paragraph()
    _document.save(f"{_token}.doc")
    return _document


async def run_command(_document: Document, _token: str, _command: str, _result: str, _exp_result: str, str_err: str, str_acc: str):
    _document.add_paragraph("Команда: " + _command)
    _document.add_paragraph("Ожидаемый Результат: " + _exp_result)
    _document.add_paragraph("Фактический результат: " + _result)
    await first_test(_document, _result, _exp_result, str_err, str_acc)
    _document.save(f"{_token}.doc")


async def first_test(_document: Document, _result: str, _exp_result: str, str_err: str, str_acc: str):
    if _result == _exp_result:
        _paragraph = _document.add_paragraph()
        _run = _paragraph.add_run(str_acc)
        _run.font.color.rgb = RGBColor(0, 255, 0)
    else:
        _paragraph = _document.add_paragraph()
        _run = _paragraph.add_run(str_err)
        _run.font.color.rgb = RGBColor(255, 0, 0)
    _document.add_paragraph()